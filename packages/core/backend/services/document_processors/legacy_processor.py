"""
Legacy Document Processor
Fallback-Implementation mit PyPDF und python-docx
"""

import logging
import time
from typing import Dict, List, Any, Tuple
import pypdf
from docx import Document as DocxDocument
import markdown
import re

from services.docling_service import DocumentChunk, ProcessedDocument

logger = logging.getLogger(__name__)


class LegacyProcessor:
    """
    Legacy Document Processor als Fallback

    Verwendet:
    - PyPDF für PDF-Verarbeitung
    - python-docx für DOCX-Verarbeitung
    - markdown für Markdown-Verarbeitung
    """

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize Legacy Processor

        Args:
            chunk_size: Maximale Anzahl Wörter pro Chunk
            chunk_overlap: Überlappung zwischen Chunks in Wörtern
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.supported_types = {
            "application/pdf": self._process_pdf,
            "application/msword": self._process_doc,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": self._process_docx,
            "text/plain": self._process_text,
            "text/markdown": self._process_markdown,
        }

        logger.info("LegacyProcessor initialized")

    async def process_document(
        self, document_id: int, file_path: str, filename: str, mime_type: str
    ) -> ProcessedDocument:
        """
        Verarbeite Dokument mit Legacy-Methoden

        Args:
            document_id: ID des Dokuments in der Datenbank
            file_path: Pfad zur Datei
            filename: Originaler Dateiname
            mime_type: MIME-Type der Datei

        Returns:
            ProcessedDocument
        """
        start_time = time.time()

        try:
            if mime_type not in self.supported_types:
                raise ValueError(f"Unsupported MIME type: {mime_type}")

            logger.info(f"Processing document with Legacy Processor: {filename}")

            # Verarbeite Dokument basierend auf MIME-Type
            processor = self.supported_types[mime_type]
            raw_text, doc_metadata = await processor(file_path)

            # Erstelle Text-Chunks
            chunks = self._create_chunks(raw_text)

            # Berechne Verarbeitungszeit
            processing_time = time.time() - start_time

            # Erweitere Metadaten
            doc_metadata["processing_method"] = "legacy"
            doc_metadata["processor_type"] = "pypdf/python-docx"

            # Erstelle ProcessedDocument
            processed_doc = ProcessedDocument(
                document_id=document_id,
                filename=filename,
                mime_type=mime_type,
                total_pages=doc_metadata.get("pages"),
                total_chunks=len(chunks),
                chunks=chunks,
                metadata=doc_metadata,
                processing_time=processing_time,
            )

            logger.info(
                f"Legacy processing completed: {filename} "
                f"({len(chunks)} chunks, {processing_time:.2f}s)"
            )

            return processed_doc

        except Exception as e:
            logger.error(f"Legacy processing failed for {filename}: {str(e)}")
            raise

    async def _process_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Verarbeite PDF-Datei mit PyPDF"""
        try:
            text_content = []
            metadata = {}

            with open(file_path, "rb") as file:
                pdf_reader = pypdf.PdfReader(file)

                # Extrahiere Metadaten
                if pdf_reader.metadata:
                    metadata.update(
                        {
                            "title": pdf_reader.metadata.get("/Title", ""),
                            "author": pdf_reader.metadata.get("/Author", ""),
                            "subject": pdf_reader.metadata.get("/Subject", ""),
                            "creator": pdf_reader.metadata.get("/Creator", ""),
                        }
                    )

                metadata["pages"] = len(pdf_reader.pages)

                # Extrahiere Text von allen Seiten
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"[Seite {page_num}]\n{page_text}")
                    except Exception as e:
                        logger.warning(
                            f"Could not extract text from page {page_num}: {str(e)}"
                        )
                        continue

            full_text = "\n\n".join(text_content)
            return full_text, metadata

        except Exception as e:
            logger.error(f"PDF processing failed: {str(e)}")
            raise

    async def _process_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Verarbeite DOCX-Datei"""
        try:
            doc = DocxDocument(file_path)

            # Extrahiere Text
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Extrahiere Metadaten
            metadata = {
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "subject": doc.core_properties.subject or "",
                "created": doc.core_properties.created.isoformat()
                if doc.core_properties.created
                else None,
                "modified": doc.core_properties.modified.isoformat()
                if doc.core_properties.modified
                else None,
                "paragraphs": len(doc.paragraphs),
            }

            full_text = "\n\n".join(text_content)
            return full_text, metadata

        except Exception as e:
            logger.error(f"DOCX processing failed: {str(e)}")
            raise

    async def _process_doc(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Verarbeite DOC-Datei (Legacy Format)"""
        try:
            # Versuche als Text zu lesen (sehr basic)
            with open(file_path, "rb") as file:
                content = file.read()
                text_content = content.decode("utf-8", errors="ignore")

            metadata = {
                "format": "DOC (Legacy)",
                "note": "Basic text extraction - consider converting to DOCX for better results",
            }

            return text_content, metadata

        except Exception as e:
            logger.error(f"DOC processing failed: {str(e)}")
            raise

    async def _process_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Verarbeite Text-Datei"""
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                content = file.read()

            lines = content.split("\n")
            words = content.split()

            metadata = {
                "lines": len(lines),
                "words": len(words),
                "characters": len(content),
                "encoding": "utf-8",
            }

            return content, metadata

        except UnicodeDecodeError:
            try:
                with open(file_path, "r", encoding="latin-1") as file:
                    content = file.read()
                metadata = {
                    "lines": len(content.split("\n")),
                    "words": len(content.split()),
                    "characters": len(content),
                    "encoding": "latin-1",
                }
                return content, metadata
            except Exception as e:
                logger.error(f"Text processing failed: {str(e)}")
                raise

    async def _process_markdown(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Verarbeite Markdown-Datei"""
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                md_content = file.read()

            # Konvertiere Markdown zu HTML und extrahiere Plain Text
            html = markdown.markdown(md_content)
            plain_text = re.sub("<[^<]+?>", "", html)

            metadata = {
                "format": "Markdown",
                "original_markdown": md_content[:500] + "..."
                if len(md_content) > 500
                else md_content,
                "html_length": len(html),
                "plain_text_length": len(plain_text),
            }

            return plain_text, metadata

        except Exception as e:
            logger.error(f"Markdown processing failed: {str(e)}")
            raise

    def _create_chunks(self, text: str) -> List[DocumentChunk]:
        """Erstelle Text-Chunks für RAG-Processing"""
        if not text or not text.strip():
            return []

        chunks = []
        words = text.split()

        if len(words) <= self.chunk_size:
            chunks.append(
                DocumentChunk(
                    content=text, chunk_index=0, metadata={"word_count": len(words)}
                )
            )
            return chunks

        start = 0
        chunk_index = 0

        while start < len(words):
            end = min(start + self.chunk_size, len(words))
            chunk_words = words[start:end]
            chunk_text = " ".join(chunk_words)

            chunks.append(
                DocumentChunk(
                    content=chunk_text,
                    chunk_index=chunk_index,
                    metadata={
                        "word_count": len(chunk_words),
                        "start_word": start,
                        "end_word": end,
                    },
                )
            )

            start = end - self.chunk_overlap
            chunk_index += 1

            if start >= len(words) - self.chunk_overlap:
                break

        return chunks
