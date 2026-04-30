"""
Tests for the structured error-code pipeline (TF-331).

The processor raises :class:`DocumentProcessingError` with a stable code,
``process_document_content`` captures that code into ``metadata.error_code``,
and the frontend renders a localised actionable message based on the code.
These tests pin both halves of that contract.
"""

import io
from unittest.mock import MagicMock

import pytest
from docx import Document as DocxDocument

from models.document import Document, DocumentStatus
from services.document_errors import (
    BINARY_CONTENT,
    DocumentProcessingError,
    EMPTY_DOCUMENT,
    LEGACY_DOC_FORMAT,
    UNKNOWN_ERROR,
    UNSUPPORTED_FORMAT,
    VECTORIZATION_FAILED,
    classify_error,
    known_codes,
)
from services.document_processors.pymupdf_processor import PyMuPDFProcessor
from services.document_service import DocumentService

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
DOC_MIME = "application/msword"


# ---------------------------------------------------------------------------
# DocumentProcessingError + classify_error
# ---------------------------------------------------------------------------


def test_error_subclasses_value_error():
    """Subclassing keeps backward compatibility with existing except ValueError."""
    err = DocumentProcessingError(LEGACY_DOC_FORMAT, "msg", filename="x.doc")
    assert isinstance(err, ValueError)
    assert err.code == LEGACY_DOC_FORMAT
    assert err.details == {"filename": "x.doc"}
    assert str(err) == "msg"


def test_classify_error_for_known_exception():
    err = DocumentProcessingError(EMPTY_DOCUMENT, "empty", filename="foo.docx")
    code, details = classify_error(err)
    assert code == EMPTY_DOCUMENT
    assert details == {"filename": "foo.docx"}


def test_classify_error_for_generic_exception_falls_back_to_unknown():
    code, details = classify_error(RuntimeError("oops"))
    assert code == UNKNOWN_ERROR
    assert details == {}


def test_known_codes_includes_all_module_constants():
    """Sanity: codes returned by known_codes() match what tests reference."""
    expected = {
        LEGACY_DOC_FORMAT,
        EMPTY_DOCUMENT,
        BINARY_CONTENT,
        UNSUPPORTED_FORMAT,
        VECTORIZATION_FAILED,
        UNKNOWN_ERROR,
    }
    assert expected.issubset(set(known_codes()))


# ---------------------------------------------------------------------------
# Processor raise sites — each must surface a specific code
# ---------------------------------------------------------------------------


@pytest.fixture
def processor():
    return PyMuPDFProcessor()


@pytest.mark.asyncio
async def test_legacy_doc_raises_legacy_doc_format_code(processor, tmp_path):
    """An OLE2 .doc file must raise with code=legacy_doc_format."""
    path = tmp_path / "old.doc"
    path.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 1024)

    with pytest.raises(DocumentProcessingError) as excinfo:
        await processor.process_document(
            document_id=1,
            file_path=str(path),
            filename="Pareto-Prinzip.doc",
            mime_type=DOC_MIME,
        )
    assert excinfo.value.code == LEGACY_DOC_FORMAT
    assert excinfo.value.details["filename"] == "Pareto-Prinzip.doc"


@pytest.mark.asyncio
async def test_empty_docx_raises_empty_document_code(processor, tmp_path):
    """A DOCX with no extractable text must raise with code=empty_document."""
    doc = DocxDocument()
    doc.add_paragraph("")  # only empty paragraphs
    buf = io.BytesIO()
    doc.save(buf)
    path = tmp_path / "empty.docx"
    path.write_bytes(buf.getvalue())

    with pytest.raises(DocumentProcessingError) as excinfo:
        await processor.process_document(
            document_id=2,
            file_path=str(path),
            filename="empty.docx",
            mime_type=DOCX_MIME,
        )
    assert excinfo.value.code == EMPTY_DOCUMENT
    assert excinfo.value.details["filename"] == "empty.docx"


@pytest.mark.asyncio
async def test_binary_renamed_as_md_raises_binary_content_code(processor, tmp_path):
    """A binary blob with .md extension must raise with code=binary_content."""
    path = tmp_path / "fake.md"
    path.write_bytes(bytes(range(256)) * 4)  # mostly control bytes

    with pytest.raises(DocumentProcessingError) as excinfo:
        await processor.process_document(
            document_id=3,
            file_path=str(path),
            filename="fake.md",
            mime_type="text/markdown",
        )
    assert excinfo.value.code == BINARY_CONTENT


@pytest.mark.asyncio
async def test_unsupported_mime_raises_unsupported_format_code(processor, tmp_path):
    """An unsupported MIME must raise with code=unsupported_format."""
    path = tmp_path / "x.weird"
    path.write_bytes(b"unknown")

    with pytest.raises(DocumentProcessingError) as excinfo:
        await processor.process_document(
            document_id=4,
            file_path=str(path),
            filename="x.weird",
            mime_type="application/x-weird",
        )
    assert excinfo.value.code == UNSUPPORTED_FORMAT
    assert excinfo.value.details["mime_type"] == "application/x-weird"


# ---------------------------------------------------------------------------
# Pipeline integration — error_code lands in metadata
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_legacy_doc_persists_error_code_in_metadata(tmp_path):
    """End-to-end: .doc upload → status=ERROR with error_code=legacy_doc_format.

    The frontend reads ``metadata.error_code`` to render the localised
    message, so this is the load-bearing assertion of the whole feature.
    """
    path = tmp_path / "Pareto-Prinzip.doc"
    path.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 1024)

    document = Document(
        filename="Pareto-Prinzip.doc",
        original_filename="Pareto-Prinzip.doc",
        file_path=str(path),
        file_size=path.stat().st_size,
        mime_type=DOC_MIME,
        status=DocumentStatus.UPLOADED,
        user_id=1,
        vector_collection="doc_test",
    )

    db = MagicMock()
    service = DocumentService()
    service.get_document_by_id = MagicMock(return_value=document)

    result = await service.process_document_content(document_id=42, db=db)

    assert result is None
    assert document.status == DocumentStatus.ERROR
    assert document.doc_metadata is not None
    assert document.doc_metadata["error_code"] == LEGACY_DOC_FORMAT
    assert document.doc_metadata["error_details"]["filename"] == "Pareto-Prinzip.doc"
    # Raw English message is still kept for logs / unknown-code fallback
    assert "Legacy .doc format" in document.doc_metadata["error"]


@pytest.mark.asyncio
async def test_empty_text_persists_empty_document_code(tmp_path):
    """End-to-end: empty .txt → metadata.error_code=empty_document."""
    path = tmp_path / "empty.txt"
    path.write_text("   \n", encoding="utf-8")

    document = Document(
        filename="empty.txt",
        original_filename="empty.txt",
        file_path=str(path),
        file_size=path.stat().st_size,
        mime_type="text/plain",
        status=DocumentStatus.UPLOADED,
        user_id=1,
        vector_collection="doc_test",
    )

    db = MagicMock()
    service = DocumentService()
    service.get_document_by_id = MagicMock(return_value=document)

    await service.process_document_content(document_id=43, db=db)

    assert document.status == DocumentStatus.ERROR
    assert document.doc_metadata["error_code"] == EMPTY_DOCUMENT


@pytest.mark.asyncio
async def test_generic_exception_falls_back_to_unknown_error(tmp_path):
    """Non-DocumentProcessingError exceptions must still carry SOME code."""
    document = Document(
        filename="missing.pdf",
        original_filename="missing.pdf",
        file_path="/tmp/nonexistent_path_for_test.pdf",
        file_size=1,
        mime_type="application/pdf",
        status=DocumentStatus.UPLOADED,
        user_id=1,
        vector_collection="doc_test",
    )

    db = MagicMock()
    service = DocumentService()
    service.get_document_by_id = MagicMock(return_value=document)

    await service.process_document_content(document_id=44, db=db)

    assert document.status == DocumentStatus.ERROR
    # File-not-found triggers a generic exception which classify_error
    # tags as unknown_error — the frontend falls back to the raw message.
    assert document.doc_metadata.get("error_code") == UNKNOWN_ERROR
