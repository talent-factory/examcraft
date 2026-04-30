"""
Tests für vollständige DOCX-Textextraktion im PyMuPDFProcessor.

Hintergrund: python-docx's `doc.paragraphs` iteriert nur Top-Level-Body-Absätze.
Tabellen, Kopf-/Fusszeilen, Textfelder und Fussnoten werden übersprungen, was
für tabellenlastige Dokumente zu 0 Chunks und stillen Vektorisierungs-Fehlern
führt.
"""

import io

import pytest
from docx import Document as DocxDocument

from services.document_processors.pymupdf_processor import PyMuPDFProcessor

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


@pytest.fixture
def processor():
    return PyMuPDFProcessor(chunk_size=1000, chunk_overlap=200)


def _save_docx(tmp_path, builder, name="doc.docx"):
    doc = DocxDocument()
    builder(doc)
    buf = io.BytesIO()
    doc.save(buf)
    path = tmp_path / name
    path.write_bytes(buf.getvalue())
    return str(path)


@pytest.mark.asyncio
async def test_docx_with_table_content_is_extracted(processor, tmp_path):
    """Table cells must be part of extracted text (currently lost)."""

    def build(doc):
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Pareto-Prinzip"
        table.cell(0, 1).text = "80/20 Regel"
        table.cell(1, 0).text = "Wirkung"
        table.cell(1, 1).text = "Ursache"

    path = _save_docx(tmp_path, build)
    result = await processor.process_document(
        document_id=1, file_path=path, filename="table.docx", mime_type=DOCX_MIME
    )

    full_text = " ".join(c.content for c in result.chunks)
    assert result.total_chunks > 0, "Expected non-zero chunks for table-only docx"
    assert "Pareto-Prinzip" in full_text
    assert "80/20 Regel" in full_text
    assert "Ursache" in full_text


@pytest.mark.asyncio
async def test_docx_header_and_footer_are_extracted(processor, tmp_path):
    """Header/Footer text must be part of extracted content."""

    def build(doc):
        doc.add_paragraph("Body content here")
        section = doc.sections[0]
        section.header.paragraphs[0].text = "HEADER: Document Title"
        section.footer.paragraphs[0].text = "FOOTER: Confidential"

    path = _save_docx(tmp_path, build)
    result = await processor.process_document(
        document_id=2,
        file_path=path,
        filename="hdrftr.docx",
        mime_type=DOCX_MIME,
    )

    full_text = " ".join(c.content for c in result.chunks)
    assert "Body content here" in full_text
    assert "HEADER: Document Title" in full_text
    assert "FOOTER: Confidential" in full_text


@pytest.mark.asyncio
async def test_docx_nested_table_content_is_extracted(processor, tmp_path):
    """Nested tables (cells containing tables) must also be walked."""

    def build(doc):
        outer = doc.add_table(rows=1, cols=1)
        cell = outer.cell(0, 0)
        cell.text = "Outer cell"
        inner = cell.add_table(rows=1, cols=1)
        inner.cell(0, 0).text = "Nested table content"

    path = _save_docx(tmp_path, build)
    result = await processor.process_document(
        document_id=3,
        file_path=path,
        filename="nested.docx",
        mime_type=DOCX_MIME,
    )

    full_text = " ".join(c.content for c in result.chunks)
    assert "Outer cell" in full_text
    assert "Nested table content" in full_text


@pytest.mark.asyncio
async def test_docx_with_mixed_content_preserves_all(processor, tmp_path):
    """Combined content (paragraphs + tables + headers) must all be extracted."""

    def build(doc):
        doc.add_paragraph("Introduction paragraph")
        doc.add_heading("Section 1", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Key fact A"
        table.cell(0, 1).text = "Key fact B"
        doc.add_paragraph("Closing paragraph")
        doc.sections[0].header.paragraphs[0].text = "Page header"

    path = _save_docx(tmp_path, build)
    result = await processor.process_document(
        document_id=4,
        file_path=path,
        filename="mixed.docx",
        mime_type=DOCX_MIME,
    )

    full_text = " ".join(c.content for c in result.chunks)
    for expected in [
        "Introduction paragraph",
        "Section 1",
        "Key fact A",
        "Key fact B",
        "Closing paragraph",
        "Page header",
    ]:
        assert expected in full_text, f"Missing in extraction: {expected!r}"


@pytest.mark.asyncio
async def test_empty_docx_raises_clear_error(processor, tmp_path):
    """Documents that yield no extractable text must fail loudly, not silently."""

    def build(doc):
        # Only empty paragraphs, no tables, no headers — no text at all
        doc.add_paragraph("")
        doc.add_paragraph("   ")

    path = _save_docx(tmp_path, build)

    with pytest.raises(ValueError, match="(?i)no.*text|empty|leer"):
        await processor.process_document(
            document_id=5,
            file_path=path,
            filename="empty.docx",
            mime_type=DOCX_MIME,
        )
