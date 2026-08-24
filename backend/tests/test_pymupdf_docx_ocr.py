"""Tests for OCR escalation of scanned DOCX files in PyMuPDFProcessor (TF-360).

Background: a "scanned" DOCX stores its pages as embedded raster images with
little/no ``<w:t>`` text. Before TF-360, ``_process_docx`` only extracted the
text layer (e.g. a 57-character source line) and also did not set ``pages``
in the metadata — leaving both escalation heuristics of the quality gate
(``scanned_low_text``, ``single_chunk_large_file``) blind. An otherwise
identical PDF, by contrast, was cleanly OCR-escalated.

These tests cover both gaps:
1. ``pages`` is set -> the quality gate can detect a scanned DOCX.
2. With OCR enabled, embedded images are recognized via Tesseract and
   included in the extracted text (analogous to the PDF escalation).

As with the PDF OCR tests, the actual Tesseract engine is mocked — the real
round trip is verified manually in the image (no Tesseract binary in CI).
"""

import io
import struct
import zlib

import pytest
from docx import Document as DocxDocument

from services.document_errors import OCR_ENGINE_FAILURE, DocumentProcessingError
from services.document_processors.pymupdf_processor import PyMuPDFProcessor

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _make_png(width: int = 8, height: int = 8, red: int = 255) -> bytes:
    """Generate a valid PNG without Pillow (stdlib zlib/struct).

    The color varies via ``red`` so different calls produce byte-distinct
    images — otherwise python-docx dedupes identical bytes into a single
    relationship.
    """

    def _chunk(typ: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + typ
            + data
            + struct.pack(">I", zlib.crc32(typ + data) & 0xFFFFFFFF)
        )

    signature = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)  # 8-bit RGB
    raw = b"".join(b"\x00" + bytes((red, 0, 0)) * width for _ in range(height))
    idat = zlib.compress(raw)
    return (
        signature + _chunk(b"IHDR", ihdr) + _chunk(b"IDAT", idat) + _chunk(b"IEND", b"")
    )


def _save_docx(tmp_path, builder, name="scanned.docx") -> str:
    doc = DocxDocument()
    builder(doc)
    buf = io.BytesIO()
    doc.save(buf)
    path = tmp_path / name
    path.write_bytes(buf.getvalue())
    return str(path)


def _scanned_docx(tmp_path, n_images: int = 1, text: str = "Quelle: Lehrmittel") -> str:
    """Build a DOCX that mimics a scanned one: little text + image(s)."""

    def build(doc):
        if text:
            doc.add_paragraph(text)
        for i in range(n_images):
            doc.add_picture(io.BytesIO(_make_png(red=50 + i * 40)))

    return _save_docx(tmp_path, build)


@pytest.mark.asyncio
async def test_docx_sets_page_count_so_quality_gate_can_fire(tmp_path):
    """Regression (cause 1): _process_docx must set pages.

    Without page_count, scanned_low_text and single_chunk_large_file were
    ineffective for DOCX and a scanned DOCX would silently pass as 'ok'.
    """
    processor = PyMuPDFProcessor()
    path = _scanned_docx(tmp_path, n_images=1)

    result = await processor.process_document(
        document_id=1, file_path=path, filename="scanned.docx", mime_type=DOCX_MIME
    )

    assert result.total_pages is not None
    assert result.total_pages >= 1


@pytest.mark.asyncio
async def test_docx_ocr_extracts_image_text_when_enabled(tmp_path, monkeypatch):
    """Cause 2: with enable_ocr=True, embedded images are OCR'd."""
    processor = PyMuPDFProcessor(enable_ocr=True)
    path = _scanned_docx(tmp_path, n_images=1)

    monkeypatch.setattr(
        processor,
        "_ocr_image_bytes",
        lambda image_bytes, ext, page_label: "GESCANNTER SEITENTEXT",
    )

    result = await processor.process_document(
        document_id=2, file_path=path, filename="scanned.docx", mime_type=DOCX_MIME
    )

    full_text = " ".join(c.content for c in result.chunks)
    assert "GESCANNTER SEITENTEXT" in full_text
    assert result.total_chunks > 0


@pytest.mark.asyncio
async def test_docx_ocr_not_run_when_disabled(tmp_path, monkeypatch):
    """First run (enable_ocr=False) must not OCR any images."""
    processor = PyMuPDFProcessor(enable_ocr=False)
    path = _scanned_docx(tmp_path, n_images=1)

    calls = []

    def _spy(image_bytes, ext, page_label):
        calls.append(page_label)
        return "SHOULD NOT APPEAR"

    monkeypatch.setattr(processor, "_ocr_image_bytes", _spy)

    result = await processor.process_document(
        document_id=3, file_path=path, filename="scanned.docx", mime_type=DOCX_MIME
    )

    full_text = " ".join(c.content for c in result.chunks)
    assert calls == [], "OCR darf im Erstlauf nicht laufen"
    assert "SHOULD NOT APPEAR" not in full_text


@pytest.mark.asyncio
async def test_docx_ocr_preserves_image_order(tmp_path, monkeypatch):
    """OCR text of multiple images must appear in document order."""
    processor = PyMuPDFProcessor(enable_ocr=True)
    path = _scanned_docx(tmp_path, n_images=3, text="")

    counter = {"n": 0}

    def _fake(image_bytes, ext, page_label):
        counter["n"] += 1
        return f"SEITE_{counter['n']}"

    monkeypatch.setattr(processor, "_ocr_image_bytes", _fake)

    result = await processor.process_document(
        document_id=4, file_path=path, filename="scanned.docx", mime_type=DOCX_MIME
    )

    full_text = "\n".join(c.content for c in result.chunks)
    pos1 = full_text.find("SEITE_1")
    pos2 = full_text.find("SEITE_2")
    pos3 = full_text.find("SEITE_3")
    assert -1 < pos1 < pos2 < pos3


@pytest.mark.asyncio
async def test_docx_ocr_engine_failure_propagates(tmp_path, monkeypatch):
    """Fatal OCR engine errors must NOT be swallowed as an empty doc."""
    processor = PyMuPDFProcessor(enable_ocr=True)
    path = _scanned_docx(tmp_path, n_images=1)

    def _boom(image_bytes, ext, page_label):
        raise DocumentProcessingError(
            OCR_ENGINE_FAILURE, "Tesseract kaputt", filename=page_label
        )

    monkeypatch.setattr(processor, "_ocr_image_bytes", _boom)

    with pytest.raises(DocumentProcessingError) as exc_info:
        await processor.process_document(
            document_id=5, file_path=path, filename="scanned.docx", mime_type=DOCX_MIME
        )
    assert exc_info.value.code == OCR_ENGINE_FAILURE


@pytest.mark.asyncio
async def test_docx_ocr_skips_unreadable_image_but_keeps_text(tmp_path, monkeypatch):
    """A single non-rasterizable image is tolerable (not fatal)."""
    processor = PyMuPDFProcessor(enable_ocr=True)
    path = _scanned_docx(tmp_path, n_images=1, text="Erhaltener Textinhalt")

    def _unsupported(image_bytes, ext, page_label):
        raise RuntimeError("kein Raster (z. B. EMF)")

    monkeypatch.setattr(processor, "_ocr_image_bytes", _unsupported)

    result = await processor.process_document(
        document_id=6, file_path=path, filename="scanned.docx", mime_type=DOCX_MIME
    )

    full_text = " ".join(c.content for c in result.chunks)
    assert "Erhaltener Textinhalt" in full_text


@pytest.mark.asyncio
async def test_docx_ocr_discarded_images_counted(tmp_path, monkeypatch):
    """A non-fatal OCR abort on one of two images is counted."""
    processor = PyMuPDFProcessor(enable_ocr=True)
    path = _scanned_docx(tmp_path, n_images=2, text="")

    counter = {"n": 0}

    def _fake(image_bytes, ext, page_label):
        counter["n"] += 1
        if counter["n"] == 1:
            raise ValueError("OOM-killed / malformed textpage")
        return "ZWEITES BILD"

    monkeypatch.setattr(processor, "_ocr_image_bytes", _fake)

    result = await processor.process_document(
        document_id=10, file_path=path, filename="scanned.docx", mime_type=DOCX_MIME
    )

    full_text = " ".join(c.content for c in result.chunks)
    assert "ZWEITES BILD" in full_text  # surviving image is retained
    assert result.metadata["ocr_pages_attempted"] == 2
    assert result.metadata["ocr_pages_discarded"] == 1
