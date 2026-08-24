"""
PyMuPDF Document Processor
Fast and efficient PDF processing using PyMuPDF (fitz)
"""

import logging
import os
import time
from typing import Dict, Iterable, List, Any, Optional, Tuple
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.oxml.ns import qn
import markdown
import re

from services.docling_service import DocumentChunk, ProcessedDocument
from .chunking import DEFAULT_MAX_CHARS_PER_CHUNK, create_chunks
from services.document_errors import (
    BINARY_CONTENT,
    EMPTY_DOCUMENT,
    LEGACY_DOC_FORMAT,
    OCR_ENGINE_FAILURE,
    UNSUPPORTED_FORMAT,
    DocumentProcessingError,
)

logger = logging.getLogger(__name__)


def _looks_like_text(content: str, threshold: float = 0.85) -> bool:
    """Return True when ``content`` looks like real text.

    Used to reject binary blobs that decoded successfully via Latin-1
    (which maps every byte 0–255 to a code point and therefore never
    raises). Real text is mostly printable plus tabs/newlines; binary
    is mostly control characters and high-bit noise.
    """
    if not content:
        return True  # empty content is handled by the upstream 0-chunk guard
    printable = sum(1 for c in content if c.isprintable() or c in "\t\n\r")
    return printable / len(content) >= threshold


def _iter_docx_text_blocks(doc) -> Iterable[str]:
    """Yield all visible text blocks in a python-docx Document.

    `doc.paragraphs` only walks the top-level body; tables, nested tables,
    headers, footers and text-frames are skipped. We iterate every `<w:t>`
    element in the document body instead, then add headers/footers from each
    section explicitly (those live outside the body XML tree).
    """

    seen_ids = set()

    body = doc.element.body
    for t_elem in body.iter(qn("w:t")):
        text = t_elem.text
        if text and text.strip():
            yield text

    for section in doc.sections:
        for hdr_ftr in (section.header, section.footer):
            for paragraph in hdr_ftr.paragraphs:
                # `Paragraph` elements may repeat across sections that share
                # a header part; dedupe on the underlying lxml element id.
                pid = id(paragraph._p)
                if pid in seen_ids:
                    continue
                seen_ids.add(pid)
                if paragraph.text.strip():
                    yield paragraph.text


def _iter_docx_image_blobs(doc) -> Iterable[Tuple[str, bytes]]:
    """Yield ``(ext, image_bytes)`` for each distinct body image, in order.

    A „scanned" DOCX stores its pages as embedded raster images. We walk
    ``<a:blip r:embed=...>`` in document order so OCR'd text keeps the visual
    reading order, and deduplicate on relationship id (Word reuses one image
    part for byte-identical pictures) so a repeated image is OCR'd only once.

    Header/footer images live on separate parts and are intentionally skipped —
    they are typically logos, not scanned content, and would add OCR noise.
    """
    body = doc.element.body
    rels = doc.part.rels
    seen: set = set()
    for blip in body.iter(qn("a:blip")):
        rid = blip.get(qn("r:embed"))
        if not rid or rid in seen:
            continue
        try:
            rel = rels[rid]
        except KeyError:
            continue
        # reltype is readable without touching target_part; external rels have
        # no local blob, so guard before dereferencing target_part.
        if rel.is_external or not rel.reltype.endswith("/image"):
            continue
        seen.add(rid)
        partname = str(rel.target_part.partname)
        ext = partname.rsplit(".", 1)[-1].lower() if "." in partname else "png"
        yield ext, rel.target_part.blob


def _read_docx_page_count(doc) -> Optional[int]:
    """Best-effort rendered page count from ``docProps/app.xml``.

    Word records the last-saved page count as ``<Pages>`` in the extended-
    properties part. python-docx does not model that part, so we read it
    straight from the package blob via a byte-level regex (namespace-agnostic,
    no XML parse needed). Returns ``None`` when absent.

    The quality gate (TF-360) gates both escalation heuristics on page_count;
    without it a scanned DOCX (page_count=0) silently passed as 'ok'.
    """
    try:
        for part in doc.part.package.iter_parts():
            if str(part.partname).endswith("/app.xml"):
                match = re.search(rb"<Pages>\s*(\d+)\s*</Pages>", part.blob)
                if match:
                    return int(match.group(1))
                break
    except Exception as exc:  # defensive: page count is best-effort, never fatal
        logger.debug(f"Could not read DOCX page count from app.xml: {exc}")
    return None


class PyMuPDFProcessor:
    """
    Modern document processor based on PyMuPDF

    Features:
    - Fast PDF processing using PyMuPDF (fitz)
    - Text extraction with layout awareness
    - Metadata extraction (author, title, creation date)
    - Multi-format support (PDF, DOCX, TXT, Markdown)
    - Optimized for performance and speed
    """

    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        enable_ocr: bool = False,
        max_chars_per_chunk: int = DEFAULT_MAX_CHARS_PER_CHUNK,
    ):
        """
        Initialize PyMuPDF Processor

        Args:
            chunk_size: Maximum number of words per chunk
            chunk_overlap: Overlap between chunks in words
            enable_ocr: If True, enables Tesseract OCR for scanned pages
                (requires an installed Tesseract + TESSDATA_PREFIX).
            max_chars_per_chunk: Generic character upper bound per chunk
                (TF-445). Over-long content is split into multiple chunks
                while preserving content, instead of being truncated at the
                embedding boundary.
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.max_chars_per_chunk = max_chars_per_chunk
        # OCR escalation (TF-360): when enabled, _process_pdf extracts the
        # text of scanned pages via Tesseract (PyMuPDF get_textpage_ocr).
        self.enable_ocr = enable_ocr
        self.ocr_language = os.getenv("OCR_LANGUAGE", "deu+eng")
        self.ocr_dpi = int(os.getenv("OCR_DPI", "200"))
        self.supported_types = {
            "application/pdf": self._process_pdf,
            "application/msword": self._process_doc,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": self._process_docx,
            "text/plain": self._process_text,
            "text/markdown": self._process_markdown,
        }

        logger.info("PyMuPDFProcessor initialized (fast PDF processing)")

    async def process_document(
        self, document_id: int, file_path: str, filename: str, mime_type: str
    ) -> ProcessedDocument:
        """
        Process document using PyMuPDF

        Args:
            document_id: ID of the document in the database
            file_path: Path to the file
            filename: Original filename
            mime_type: MIME type of the file

        Returns:
            ProcessedDocument with extended metadata
        """
        start_time = time.time()

        try:
            if mime_type not in self.supported_types:
                raise DocumentProcessingError(
                    UNSUPPORTED_FORMAT,
                    f"Unsupported MIME type: {mime_type}",
                    filename=filename,
                    mime_type=mime_type,
                )

            ocr_state = "OCR enabled" if self.enable_ocr else "no OCR"
            logger.info(f"Processing document with PyMuPDF ({ocr_state}): {filename}")

            # Process document based on MIME type
            processor = self.supported_types[mime_type]
            raw_text, doc_metadata = await processor(file_path, filename)

            # Create text chunks
            chunks = self._create_chunks(raw_text)

            # Calculate processing time
            processing_time = time.time() - start_time

            # Extend metadata
            doc_metadata["processing_method"] = "pymupdf"
            doc_metadata["processor_type"] = "PyMuPDF"
            doc_metadata["ocr_enabled"] = self.enable_ocr

            # Build ProcessedDocument
            processed_doc = ProcessedDocument(
                document_id=document_id,
                filename=filename,
                mime_type=mime_type,
                total_pages=doc_metadata.get("pages"),
                total_chunks=len(chunks),
                chunks=chunks,
                metadata=doc_metadata,
                processing_time=processing_time,
            )

            logger.info(
                f"PyMuPDF processing completed ({ocr_state}): {filename} "
                f"({len(chunks)} chunks, {processing_time:.2f}s)"
            )

            return processed_doc

        except Exception as e:
            logger.error(f"PyMuPDF processing failed for {filename}: {str(e)}")
            raise

    async def _process_pdf(
        self, file_path: str, filename: str
    ) -> Tuple[str, Dict[str, Any]]:
        """
        Process PDF file using PyMuPDF (very fast)

        Args:
            file_path: Path to the PDF file
            filename: Original filename used as a fallback title

        Returns:
            Tuple (full_text, metadata)
        """
        try:
            text_content = []
            metadata = {}

            # Open PDF using PyMuPDF
            doc = fitz.open(file_path)

            # Extract metadata
            pdf_metadata = doc.metadata
            if pdf_metadata:
                # Title (with fallback to filename)
                title = pdf_metadata.get("title", "").strip()
                if not title:
                    title = filename.rsplit(".", 1)[0] if "." in filename else filename
                metadata["title"] = title

                # Author
                author = pdf_metadata.get("author", "").strip()
                metadata["author"] = author

                # Subject
                subject = pdf_metadata.get("subject", "").strip()
                if subject:
                    metadata["subject"] = subject

                # Keywords
                keywords = pdf_metadata.get("keywords", "").strip()
                if keywords:
                    metadata["keywords"] = keywords

                # Creation Date
                creation_date = pdf_metadata.get("creationDate", "").strip()
                if creation_date:
                    # PyMuPDF format: D:20240101120000+01'00'
                    try:
                        # Strip the 'D:' prefix
                        date_str = creation_date.replace("D:", "")
                        # Extract YYYYMMDD
                        if len(date_str) >= 8:
                            year = date_str[0:4]
                            month = date_str[4:6]
                            day = date_str[6:8]
                            metadata["creation_date"] = f"{year}-{month}-{day}"
                    except Exception as e:
                        logger.debug(f"Failed to parse PDF creation date: {e}")

            # Page count
            metadata["pages"] = doc.page_count

            # Extract text from all pages
            ocr_pages_attempted = 0
            ocr_pages_discarded = 0
            for page_num in range(doc.page_count):
                try:
                    page = doc[page_num]
                    if self.enable_ocr:
                        # full=False -> Tesseract only runs on pages without
                        # a text layer (scanned pages); existing text is kept.
                        ocr_pages_attempted += 1
                        try:
                            page_text = self._ocr_pdf_page(page, page_num, filename)
                        except DocumentProcessingError:
                            # Propagate fatal engine errors loudly.
                            raise
                        except Exception as ocr_page_err:
                            # Non-fatal per-page OCR failure (OOM-killed
                            # subprocess, malformed textpage). Count it so the
                            # quality gate sees the gap (TF-367), but skip the
                            # page instead of losing the whole document.
                            ocr_pages_discarded += 1
                            logger.warning(
                                f"OCR auf Seite {page_num + 1} verworfen "
                                f"(nicht-fataler Abbruch): {ocr_page_err}"
                            )
                            continue
                    else:
                        page_text = page.get_text("text")  # Plain text extraction

                    if page_text.strip():
                        text_content.append(f"[Seite {page_num + 1}]\n{page_text}")
                except DocumentProcessingError:
                    # Propagate fatal OCR errors loudly, don't skip them.
                    raise
                except Exception as e:
                    # Page-local decode errors (PDF corruption) are tolerable.
                    logger.warning(
                        f"Could not extract text from page {page_num + 1}: {str(e)}"
                    )
                    continue

            # Close PDF
            doc.close()

            # Surface the OCR discard counters for the quality gate (TF-367).
            if self.enable_ocr:
                metadata["ocr_pages_attempted"] = ocr_pages_attempted
                metadata["ocr_pages_discarded"] = ocr_pages_discarded

            # Extract headings from text (simple heuristic)
            full_text = "\n\n".join(text_content)
            sections = self._extract_headings_from_text(full_text)
            if sections:
                metadata["sections"] = sections[:30]  # Max 30 headings
                metadata["section_count"] = len(sections[:30])

            # Defaults if metadata is missing
            if "title" not in metadata or not metadata["title"]:
                metadata["title"] = (
                    filename.rsplit(".", 1)[0] if "." in filename else filename
                )
            if "author" not in metadata:
                metadata["author"] = ""

            return full_text, metadata

        except Exception as e:
            logger.error(f"PDF processing failed: {str(e)}")
            raise

    def _extract_headings_from_text(self, text: str) -> List[str]:
        """
        Extract potential headings from text

        Heuristic:
        - Lines with <= 100 characters
        - Doesn't end with a period
        - Starts with an uppercase letter or a digit
        - Has at least 3 words

        Args:
            text: Full text

        Returns:
            List of headings
        """
        headings = []
        lines = text.split("\n")

        for line in lines:
            line = line.strip()

            # Filter out lines that are too long
            if len(line) > 100 or len(line) < 10:
                continue

            # Filter out lines that end with a period (normal sentences)
            if line.endswith("."):
                continue

            # Must start with an uppercase letter or a digit
            if not (line[0].isupper() or line[0].isdigit()):
                continue

            # Must have at least 3 words
            words = line.split()
            if len(words) < 3:
                continue

            # Filter out lines with lots of special characters
            special_char_count = sum(1 for c in line if c in "()[]{}|\\/@#$%^&*+=~`")
            if special_char_count > 3:
                continue

            headings.append(line)

        return headings

    def _ocr_image_bytes(self, image_bytes: bytes, ext: str, page_label: str) -> str:
        """OCR a single embedded raster image via PyMuPDF + Tesseract.

        Opens the image as a single-page PyMuPDF document and OCRs the whole
        page (``full=True`` — an embedded image has no text layer that needs
        to be preserved). Engine errors are marked as fatal.
        """
        image_doc = fitz.open(stream=image_bytes, filetype=ext)
        try:
            page = image_doc[0]
            try:
                ocr_tp = page.get_textpage_ocr(
                    language=self.ocr_language, dpi=self.ocr_dpi, full=True
                )
            except RuntimeError as ocr_err:
                # An engine error (Tesseract missing/misconfigured, language
                # pack missing) is fatal for the whole document — propagate it
                # loudly instead of silently swallowing it as an empty doc
                # (TF-360).
                raise DocumentProcessingError(
                    OCR_ENGINE_FAILURE,
                    f"Tesseract-OCR fehlgeschlagen für {page_label}: {ocr_err}",
                    filename=page_label,
                ) from ocr_err
            return page.get_text("text", textpage=ocr_tp)
        finally:
            image_doc.close()

    def _ocr_pdf_page(self, page, page_num: int, filename: str) -> str:
        """OCR a single PDF page via PyMuPDF/Tesseract.

        Engine errors (``RuntimeError``: Tesseract missing/misconfigured,
        language pack missing) are fatal for the whole document and are
        reported as ``DocumentProcessingError(OCR_ENGINE_FAILURE)``. All
        other exceptions (OOM-killed subprocess, malformed textpage)
        propagate raw — the calling loop counts them as a discarded page
        (TF-367) and skips it, instead of failing the whole document.
        """
        try:
            ocr_tp = page.get_textpage_ocr(
                language=self.ocr_language, dpi=self.ocr_dpi, full=False
            )
        except RuntimeError as ocr_err:
            raise DocumentProcessingError(
                OCR_ENGINE_FAILURE,
                f"Tesseract-OCR fehlgeschlagen auf Seite {page_num + 1}: {ocr_err}",
                filename=filename,
            ) from ocr_err
        return page.get_text("text", textpage=ocr_tp)

    async def _process_docx(
        self, file_path: str, filename: str
    ) -> Tuple[str, Dict[str, Any]]:
        """Process DOCX file (body + tables + header/footer + OCR).

        When OCR is enabled (escalation, TF-360), embedded images are also
        recognized via Tesseract — this makes a "scanned" DOCX (pages as
        images, barely any ``<w:t>`` text) usable, analogous to the PDF
        escalation. The first pass (``enable_ocr=False``) only extracts the
        text layer; the quality gate flags the thin result (see ``pages``
        below) and the reprocess comes back here with ``enable_ocr=True``.
        """
        try:
            doc = DocxDocument(file_path)

            text_content = list(_iter_docx_text_blocks(doc))

            # Collect embedded images (order preserved, deduplicated). Always
            # count them — the count serves as a page fallback for the
            # quality gate; OCR only runs when escalation is enabled.
            image_blobs = list(_iter_docx_image_blobs(doc))

            ocr_pages_attempted = 0
            ocr_pages_discarded = 0
            if self.enable_ocr and image_blobs:
                for idx, (ext, blob) in enumerate(image_blobs, start=1):
                    ocr_pages_attempted += 1
                    label = f"{filename}#Bild{idx}"
                    try:
                        ocr_text = self._ocr_image_bytes(blob, ext, label)
                    except DocumentProcessingError:
                        # Propagate fatal OCR engine errors loudly.
                        raise
                    except Exception as img_err:
                        # Non-fatal per-image OCR failure (non-rasterizable
                        # image, OOM, malformed textpage): count it (TF-367)
                        # and skip it.
                        ocr_pages_discarded += 1
                        logger.warning(f"OCR übersprungen für {label}: {img_err}")
                        continue
                    if ocr_text.strip():
                        text_content.append(f"[Bild {idx}]\n{ocr_text.strip()}")

            title = doc.core_properties.title or ""
            if not title:
                title = filename.rsplit(".", 1)[0] if "." in filename else filename

            # Page count for the quality gate: real rendered page count from
            # app.xml, else the number of embedded images (scan proxy:
            # ~1/page). Without page_count, both scanned_low_text and
            # single_chunk_large_file would be ineffective for DOCX (TF-360).
            page_count = _read_docx_page_count(doc)
            if not page_count and image_blobs:
                page_count = len(image_blobs)

            metadata = {
                "title": title,
                "author": doc.core_properties.author or "",
                "subject": doc.core_properties.subject or "",
                "created": doc.core_properties.created.isoformat()
                if doc.core_properties.created
                else None,
                "modified": doc.core_properties.modified.isoformat()
                if doc.core_properties.modified
                else None,
                "paragraphs": len(doc.paragraphs),
                "text_blocks": len(text_content),
                "image_count": len(image_blobs),
            }
            if page_count:
                metadata["pages"] = page_count

            # Surface the OCR discard counters for the quality gate (TF-367).
            if self.enable_ocr:
                metadata["ocr_pages_attempted"] = ocr_pages_attempted
                metadata["ocr_pages_discarded"] = ocr_pages_discarded

            full_text = "\n\n".join(text_content)

            if not full_text.strip():
                raise DocumentProcessingError(
                    EMPTY_DOCUMENT,
                    f"No extractable text found in DOCX '{filename}' "
                    "(document is empty or contains only images/objects)",
                    filename=filename,
                )

            return full_text, metadata

        except ValueError:
            raise
        except Exception as e:
            logger.error(f"DOCX processing failed: {str(e)}")
            raise

    async def _process_doc(
        self, file_path: str, filename: str
    ) -> Tuple[str, Dict[str, Any]]:
        """Process DOC file (legacy CFB/OLE2 format).

        `.doc` is a binary OLE2 compound file. Without `antiword`,
        `libreoffice --headless` or a CFB parser, we cannot reliably extract
        text. The previous implementation decoded raw bytes as UTF-8 with
        `errors='ignore'`, producing garbage that silently embedded as
        nonsense vectors. We now refuse the file with a clear error so the
        document is marked ERROR and the user knows to convert to DOCX.
        """
        with open(file_path, "rb") as fh:
            header = fh.read(8)

        if header.startswith(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"):
            raise DocumentProcessingError(
                LEGACY_DOC_FORMAT,
                f"Legacy .doc format detected for '{filename}'. "
                "Reliable text extraction requires conversion. "
                "Please save the document as .docx and re-upload.",
                filename=filename,
            )

        # Some `.doc` files in the wild are actually mislabeled RTF or text.
        # Try a best-effort decode and refuse if no meaningful text remains.
        with open(file_path, "rb") as fh:
            raw = fh.read()
        text_content = raw.decode("utf-8", errors="ignore")
        # Strip control bytes that survived decode
        cleaned = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text_content)
        if len(cleaned.strip()) < 50:
            raise DocumentProcessingError(
                EMPTY_DOCUMENT,
                f"No extractable text found in legacy .doc '{filename}'. "
                "Please save the document as .docx and re-upload.",
                filename=filename,
            )

        metadata = {
            "title": filename.rsplit(".", 1)[0] if "." in filename else filename,
            "format": "DOC (Legacy)",
            "note": "Best-effort text extraction; converting to DOCX is recommended.",
        }
        return cleaned, metadata

    async def _process_text(
        self, file_path: str, filename: str
    ) -> Tuple[str, Dict[str, Any]]:
        """Process text file with encoding fallback (UTF-8 → Latin-1)."""
        try:
            content, encoding = self._read_text_with_fallback(file_path)
        except Exception as e:
            logger.error(f"Text processing failed: {str(e)}")
            raise

        metadata = {
            "title": filename.rsplit(".", 1)[0] if "." in filename else filename,
            "lines": len(content.split("\n")),
            "words": len(content.split()),
            "characters": len(content),
            "encoding": encoding,
        }
        return content, metadata

    async def _process_markdown(
        self, file_path: str, filename: str
    ) -> Tuple[str, Dict[str, Any]]:
        """Process Markdown file with encoding fallback (UTF-8 → Latin-1)."""
        try:
            md_content, encoding = self._read_text_with_fallback(file_path)

            # Convert Markdown to HTML and extract plain text
            html = markdown.markdown(md_content)
            plain_text = re.sub("<[^<]+?>", "", html)

            # Extract headings from Markdown
            sections = []
            heading_pattern = r"^(#+)\s+(.+)$"
            matches = re.finditer(heading_pattern, md_content, re.MULTILINE)

            for match in matches:
                title = match.group(2).strip()
                if len(title) > 200:  # Filter out headings that are too long
                    continue
                sections.append(title)

            metadata = {
                "title": filename.rsplit(".", 1)[0] if "." in filename else filename,
                "format": "Markdown",
                "encoding": encoding,
                "sections": sections[:30] if sections else [],
                "section_count": len(sections[:30]) if sections else 0,
                "html_length": len(html),
                "plain_text_length": len(plain_text),
            }

            return plain_text, metadata

        except Exception as e:
            logger.error(f"Markdown processing failed: {str(e)}")
            raise

    @staticmethod
    def _read_text_with_fallback(file_path: str) -> Tuple[str, str]:
        """Read a text file as UTF-8, falling back to Latin-1.

        Returns ``(content, encoding)`` so callers can record the encoding
        in metadata for diagnostics. Latin-1 decodes any byte sequence —
        including binary files renamed with a text extension — so when
        the fallback fires we additionally check the printable-character
        ratio and refuse mojibake. Failures of the fallback ``open()``
        itself are wrapped to preserve the file path in the error.
        """
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                return file.read(), "utf-8"
        except UnicodeDecodeError:
            pass

        try:
            with open(file_path, "r", encoding="latin-1") as file:
                content = file.read()
        except OSError as e:
            raise OSError(f"Latin-1 fallback read failed for {file_path}: {e}") from e

        if not _looks_like_text(content):
            raise DocumentProcessingError(
                BINARY_CONTENT,
                f"File {file_path} does not appear to be a text file "
                "(low printable-character ratio after Latin-1 fallback). "
                "Likely binary content with a misleading extension.",
            )

        logger.warning(f"File {file_path} is not valid UTF-8; decoded as Latin-1")
        return content, "latin-1"

    def _create_chunks(self, text: str) -> List[DocumentChunk]:
        """Create text chunks for RAG processing"""
        return create_chunks(
            text,
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            max_chars=self.max_chars_per_chunk,
        )
