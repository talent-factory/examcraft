"""
Legacy Document Processor
Fallback implementation using PyPDF and python-docx
"""

import logging
import time
from typing import Dict, List, Any, Tuple
import pypdf
from docx import Document as DocxDocument
import markdown
import re

from services.docling_service import DocumentChunk, ProcessedDocument
from .chunking import DEFAULT_MAX_CHARS_PER_CHUNK, create_chunks
from services.document_errors import (
    EMPTY_DOCUMENT,
    UNSUPPORTED_FORMAT,
    DocumentProcessingError,
)

logger = logging.getLogger(__name__)


class LegacyProcessor:
    """
    Legacy Document Processor used as fallback

    Uses:
    - PyPDF for PDF processing
    - python-docx for DOCX processing
    - markdown for Markdown processing
    """

    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        max_chars_per_chunk: int = DEFAULT_MAX_CHARS_PER_CHUNK,
    ):
        """
        Initialize Legacy Processor

        Args:
            chunk_size: Maximum number of words per chunk
            chunk_overlap: Overlap between chunks in words
            max_chars_per_chunk: Generic character upper bound per chunk
                (TF-445). Over-long content is split into multiple chunks
                while preserving content, instead of being truncated at the
                embedding boundary.
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.max_chars_per_chunk = max_chars_per_chunk
        self.supported_types = {
            "application/pdf": self._process_pdf,
            "application/msword": self._process_doc,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": self._process_docx,
            "text/plain": self._process_text,
            "text/markdown": self._process_markdown,
        }

        logger.info("LegacyProcessor initialized")

    async def process_document(
        self, document_id: int, file_path: str, filename: str, mime_type: str
    ) -> ProcessedDocument:
        """
        Process document using legacy methods

        Args:
            document_id: ID of the document in the database
            file_path: Path to the file
            filename: Original filename
            mime_type: MIME type of the file

        Returns:
            ProcessedDocument
        """
        start_time = time.time()

        try:
            if mime_type not in self.supported_types:
                raise DocumentProcessingError(
                    UNSUPPORTED_FORMAT,
                    f"Unsupported MIME type: {mime_type}",
                    filename=filename,
                    mime_type=mime_type,
                )

            logger.info(f"Processing document with Legacy Processor: {filename}")

            # Process document based on MIME type
            processor = self.supported_types[mime_type]
            raw_text, doc_metadata = await processor(file_path)

            # Create text chunks
            chunks = self._create_chunks(raw_text)

            # Calculate processing time
            processing_time = time.time() - start_time

            # Extend metadata
            doc_metadata["processing_method"] = "legacy"
            doc_metadata["processor_type"] = "pypdf/python-docx"

            # Build ProcessedDocument
            processed_doc = ProcessedDocument(
                document_id=document_id,
                filename=filename,
                mime_type=mime_type,
                total_pages=doc_metadata.get("pages"),
                total_chunks=len(chunks),
                chunks=chunks,
                metadata=doc_metadata,
                processing_time=processing_time,
            )

            logger.info(
                f"Legacy processing completed: {filename} "
                f"({len(chunks)} chunks, {processing_time:.2f}s)"
            )

            return processed_doc

        except Exception as e:
            logger.error(f"Legacy processing failed for {filename}: {str(e)}")
            raise

    async def _process_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process PDF file using PyPDF"""
        try:
            text_content = []
            metadata = {}

            with open(file_path, "rb") as file:
                pdf_reader = pypdf.PdfReader(file)

                # Extract metadata
                if pdf_reader.metadata:
                    metadata.update(
                        {
                            "title": pdf_reader.metadata.get("/Title", ""),
                            "author": pdf_reader.metadata.get("/Author", ""),
                            "subject": pdf_reader.metadata.get("/Subject", ""),
                            "creator": pdf_reader.metadata.get("/Creator", ""),
                        }
                    )

                metadata["pages"] = len(pdf_reader.pages)

                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"[Seite {page_num}]\n{page_text}")
                    except Exception as e:
                        logger.warning(
                            f"Could not extract text from page {page_num}: {str(e)}"
                        )
                        continue

            full_text = "\n\n".join(text_content)
            return full_text, metadata

        except Exception as e:
            logger.error(f"PDF processing failed: {str(e)}")
            raise

    async def _process_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process DOCX file (body + tables + header/footer)."""
        from services.document_processors.pymupdf_processor import (
            _iter_docx_text_blocks,
        )

        try:
            doc = DocxDocument(file_path)

            text_content = list(_iter_docx_text_blocks(doc))

            metadata = {
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "subject": doc.core_properties.subject or "",
                "created": doc.core_properties.created.isoformat()
                if doc.core_properties.created
                else None,
                "modified": doc.core_properties.modified.isoformat()
                if doc.core_properties.modified
                else None,
                "paragraphs": len(doc.paragraphs),
                "text_blocks": len(text_content),
            }

            full_text = "\n\n".join(text_content)
            if not full_text.strip():
                raise DocumentProcessingError(
                    EMPTY_DOCUMENT,
                    "No extractable text found in DOCX (empty or media-only)",
                )
            return full_text, metadata

        except ValueError:
            raise
        except Exception as e:
            logger.error(f"DOCX processing failed: {str(e)}")
            raise

    async def _process_doc(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process DOC file (legacy format)"""
        try:
            # Try reading as text (very basic)
            with open(file_path, "rb") as file:
                content = file.read()
                text_content = content.decode("utf-8", errors="ignore")

            metadata = {
                "format": "DOC (Legacy)",
                "note": "Basic text extraction - consider converting to DOCX for better results",
            }

            return text_content, metadata

        except Exception as e:
            logger.error(f"DOC processing failed: {str(e)}")
            raise

    async def _process_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process text file with encoding fallback (UTF-8 → Latin-1)."""
        try:
            content, encoding = self._read_text_with_fallback(file_path)
        except Exception as e:
            logger.error(f"Text processing failed: {str(e)}")
            raise

        metadata = {
            "lines": len(content.split("\n")),
            "words": len(content.split()),
            "characters": len(content),
            "encoding": encoding,
        }
        return content, metadata

    async def _process_markdown(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process Markdown file with encoding fallback (UTF-8 → Latin-1)."""
        try:
            md_content, encoding = self._read_text_with_fallback(file_path)

            # Convert Markdown to HTML and extract plain text
            html = markdown.markdown(md_content)
            plain_text = re.sub("<[^<]+?>", "", html)

            metadata = {
                "format": "Markdown",
                "encoding": encoding,
                "original_markdown": md_content[:500] + "..."
                if len(md_content) > 500
                else md_content,
                "html_length": len(html),
                "plain_text_length": len(plain_text),
            }

            return plain_text, metadata

        except Exception as e:
            logger.error(f"Markdown processing failed: {str(e)}")
            raise

    @staticmethod
    def _read_text_with_fallback(file_path: str) -> Tuple[str, str]:
        """Read a text file as UTF-8, falling back to Latin-1.

        Delegates to the same hardened helper used by ``PyMuPDFProcessor``
        so both processors share the mojibake check, the OSError context
        wrapping, and the Latin-1 warning log.
        """
        from services.document_processors.pymupdf_processor import (
            PyMuPDFProcessor,
        )

        return PyMuPDFProcessor._read_text_with_fallback(file_path)

    def _create_chunks(self, text: str) -> List[DocumentChunk]:
        """Create text chunks for RAG processing"""
        return create_chunks(
            text,
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            max_chars=self.max_chars_per_chunk,
        )
